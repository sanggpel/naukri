"""FastAPI web dashboard for tracking job applications."""

import asyncio
import os

from fastapi import FastAPI, Form, Request
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates

from ..models import Application, NetworkMatch
from ..tracker import (
    delete_application,
    get_application,
    load_applications,
    save_application,
    update_notes,
    update_referrals,
    update_status,
)
from ..network.linkedin import find_connections_at_company
from ..profile_loader import load_profile

BASE_DIR = os.path.join(os.path.dirname(__file__), "..", "..")
TEMPLATES_DIR = os.path.join(BASE_DIR, "templates", "web")
STATIC_DIR = os.path.join(BASE_DIR, "static")

templates = Jinja2Templates(directory=TEMPLATES_DIR)


def create_app() -> FastAPI:
    app = FastAPI(title="Job Application Tracker")

    app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")

    @app.get("/", response_class=HTMLResponse)
    async def dashboard(
        request: Request,
        status: str = "all",
        q: str = "",
        company: str = "",
        source: str = "",
        remote: str = "",
        scout_added: str = "",
    ):
        all_apps = load_applications()

        # Check if profile exists
        no_profile = not os.path.exists(os.path.join(BASE_DIR, "config", "profile.yaml"))

        # Build filter options from all data
        companies = sorted({a.company for a in all_apps if a.company})
        sources = sorted({a.source for a in all_apps if a.source})

        # Stats (before filtering)
        statuses = ["all", "discovered", "generated", "applied", "interviewing", "offered", "rejected", "not_relevant", "withdrawn"]
        stats = {}
        for s in statuses[1:]:
            stats[s] = sum(1 for a in all_apps if a.status == s)
        stats["all"] = len(all_apps)

        # Apply filters
        apps = list(all_apps)
        if status != "all":
            apps = [a for a in apps if a.status == status]
        if q:
            q_lower = q.lower()
            apps = [a for a in apps if q_lower in a.job_title.lower() or q_lower in a.company.lower() or q_lower in a.description.lower()]
        if company:
            apps = [a for a in apps if a.company == company]
        if source:
            apps = [a for a in apps if a.source == source]
        if remote == "yes":
            apps = [a for a in apps if any(kw in (a.location + " " + a.job_title).lower() for kw in ["remote", "hybrid", "anywhere", "work from home"])]

        # Sort by date descending
        apps.sort(key=lambda a: a.date_generated, reverse=True)

        return templates.TemplateResponse("dashboard.html", {
            "request": request,
            "applications": apps,
            "current_status": status,
            "statuses": statuses,
            "stats": stats,
            "companies": companies,
            "sources": sources,
            "filters": {"q": q, "company": company, "source": source, "remote": remote},
            "scout_added": int(scout_added) if scout_added.isdigit() else None,
            "no_profile": no_profile,
        })

    @app.get("/application/{app_id}", response_class=HTMLResponse)
    async def detail(request: Request, app_id: str, error: str = ""):
        app_data = get_application(app_id)
        if not app_data:
            return RedirectResponse("/", status_code=302)

        statuses = ["discovered", "generated", "applied", "interviewing", "offered", "rejected", "not_relevant", "withdrawn"]
        return templates.TemplateResponse("detail.html", {
            "request": request,
            "app": app_data,
            "statuses": statuses,
            "error": error,
        })

    @app.post("/application/{app_id}/status")
    async def update_app_status(app_id: str, status: str = Form(...)):
        update_status(app_id, status)
        return RedirectResponse(f"/application/{app_id}", status_code=302)

    @app.post("/bulk-status")
    async def bulk_update_status(request: Request):
        form = await request.form()
        new_status = form.get("bulk_status", "not_relevant")
        selected = form.getlist("selected")
        for app_id in selected:
            update_status(app_id, new_status)
        return RedirectResponse("/", status_code=302)

    @app.post("/application/{app_id}/notes")
    async def update_app_notes(app_id: str, notes: str = Form(...)):
        update_notes(app_id, notes)
        return RedirectResponse(f"/application/{app_id}", status_code=302)

    @app.post("/application/{app_id}/delete")
    async def delete_app(app_id: str):
        delete_application(app_id)
        return RedirectResponse("/", status_code=302)

    @app.post("/application/{app_id}/generate")
    async def generate_docs(app_id: str):
        """Generate resume + cover letter + fit summary + gap analysis for a tracked application."""
        app_data = get_application(app_id)
        if not app_data or not app_data.description:
            return RedirectResponse(f"/application/{app_id}", status_code=302)

        def _do_generate():
            from ..generator.cache import save_cover_letter_to_cache, save_resume_to_cache
            from ..generator.renderer import render_cover_letter_pdf, render_resume_pdf
            from ..generator.unified import generate_application
            from ..profile_loader import load_profile

            profile = load_profile()
            result = generate_application(profile, app_data.description)

            keywords = result["keywords"]
            resume = result["resume"]
            cover_letter = result["cover_letter"]

            job_title = app_data.job_title or keywords.job_title
            company = app_data.company or keywords.company_name

            resume_path = render_resume_pdf(resume, profile.name)
            save_resume_to_cache(keywords, resume, resume_path, job_title, company)

            cl_path = render_cover_letter_pdf(cover_letter, profile.name)
            save_cover_letter_to_cache(keywords, cover_letter, cl_path, job_title, company)

            app_data.resume_path = os.path.abspath(resume_path)
            app_data.cover_letter_path = os.path.abspath(cl_path)
            app_data.ats_score = resume.ats_score_estimate
            app_data.fit_summary = result["fit_summary"]
            app_data.gap_analysis = result["gap_analysis"]
            if app_data.status == "discovered":
                app_data.status = "generated"
            save_application(app_data)

        try:
            await asyncio.to_thread(_do_generate)
        except Exception as exc:
            import logging
            import urllib.parse
            logging.getLogger(__name__).error("Generation failed for %s: %s", app_id, exc)
            error_msg = urllib.parse.quote(str(exc)[:300])
            return RedirectResponse(f"/application/{app_id}?error={error_msg}", status_code=302)

        return RedirectResponse(f"/application/{app_id}", status_code=302)

    @app.post("/application/{app_id}/referrals")
    async def fetch_referrals(request: Request, app_id: str):
        app_data = get_application(app_id)
        error_msg = None
        if app_data and app_data.company:
            try:
                matches = await asyncio.to_thread(find_connections_at_company, app_data.company)
                update_referrals(app_id, matches)
                if not matches:
                    error_msg = f"No connections found at {app_data.company} in your LinkedIn network."
            except Exception as e:
                error_msg = str(e)

        if error_msg:
            # Re-render detail page with the error shown
            app_data = get_application(app_id)
            statuses = ["discovered", "generated", "applied", "interviewing", "offered", "rejected", "not_relevant", "withdrawn"]
            return templates.TemplateResponse("detail.html", {
                "request": request,
                "app": app_data,
                "statuses": statuses,
                "referral_error": error_msg,
            })
        return RedirectResponse(f"/application/{app_id}", status_code=302)

    @app.post("/application/{app_id}/fetch-description")
    async def fetch_desc(request: Request, app_id: str):
        """Fetch job description from the job URL using smart scraping."""
        app_data = get_application(app_id)
        if not app_data or not app_data.url:
            return RedirectResponse(f"/application/{app_id}", status_code=302)

        from ..discovery.fetcher import fetch_description

        result = await asyncio.to_thread(fetch_description, app_data.url)

        if result["description"]:
            app_data.description = result["description"]
            save_application(app_data)

        if result["error"]:
            # Re-render with error
            statuses = ["discovered", "generated", "applied", "interviewing", "offered", "rejected", "not_relevant", "withdrawn"]
            return templates.TemplateResponse("detail.html", {
                "request": request,
                "app": app_data,
                "statuses": statuses,
                "fetch_error": result["error"],
                "fetch_method": result["method"],
            })

        return RedirectResponse(f"/application/{app_id}", status_code=302)

    @app.post("/application/{app_id}/save-description")
    async def save_desc(app_id: str, description: str = Form(...)):
        """Save a manually pasted job description."""
        app_data = get_application(app_id)
        if app_data:
            app_data.description = description.strip()
            save_application(app_data)
        return RedirectResponse(f"/application/{app_id}", status_code=302)

    @app.post("/application/{app_id}/fill-gaps")
    async def fill_gaps(app_id: str, additional_context: str = Form(...)):
        """Update profile with additional context, then regenerate resume & cover letter."""
        app_data = get_application(app_id)
        if not app_data or not app_data.description or not additional_context.strip():
            return RedirectResponse(f"/application/{app_id}", status_code=302)

        def _do_fill_gaps():
            from ..generator.cache import save_cover_letter_to_cache, save_resume_to_cache
            from ..generator.renderer import render_cover_letter_pdf, render_resume_pdf
            from ..generator.unified import generate_application
            from ..profile_loader import load_profile
            from ..profile_updater import update_profile_from_context

            # Step 1: Update profile.yaml with the additional context
            update_profile_from_context(additional_context.strip(), app_data.gap_analysis)

            # Step 2: Reload the updated profile
            profile = load_profile()

            # Step 3: Regenerate with enriched profile
            result = generate_application(profile, app_data.description)

            keywords = result["keywords"]
            resume = result["resume"]
            cover_letter = result["cover_letter"]

            job_title = app_data.job_title or keywords.job_title
            company = app_data.company or keywords.company_name

            resume_path = render_resume_pdf(resume, profile.name)
            save_resume_to_cache(keywords, resume, resume_path, job_title, company)

            cl_path = render_cover_letter_pdf(cover_letter, profile.name)
            save_cover_letter_to_cache(keywords, cover_letter, cl_path, job_title, company)

            app_data.resume_path = os.path.abspath(resume_path)
            app_data.cover_letter_path = os.path.abspath(cl_path)
            app_data.ats_score = resume.ats_score_estimate
            app_data.fit_summary = result["fit_summary"]
            app_data.gap_analysis = result["gap_analysis"]
            if app_data.status == "discovered":
                app_data.status = "generated"
            save_application(app_data)

        try:
            await asyncio.to_thread(_do_fill_gaps)
        except Exception as exc:
            import logging
            import urllib.parse
            logging.getLogger(__name__).error("Fill-gaps failed for %s: %s", app_id, exc)
            error_msg = urllib.parse.quote(str(exc)[:300])
            return RedirectResponse(f"/application/{app_id}?error={error_msg}", status_code=302)

        return RedirectResponse(f"/application/{app_id}", status_code=302)

    @app.post("/add-job")
    async def add_job(
        job_title: str = Form(...),
        company: str = Form(""),
        url: str = Form(""),
        location: str = Form(""),
        description: str = Form(""),
    ):
        """Manually add a job to the tracker."""
        import hashlib
        from datetime import datetime

        job_id = hashlib.md5(
            f"{job_title}{company}{url}{datetime.now().isoformat()}".encode()
        ).hexdigest()[:12]

        app_record = Application(
            id=job_id,
            job_title=job_title.strip(),
            company=company.strip(),
            location=location.strip(),
            url=url.strip(),
            source="manual",
            date_generated=datetime.now().isoformat(),
            status="discovered",
            description=description.strip(),
        )
        save_application(app_record)
        return RedirectResponse(f"/application/{job_id}", status_code=302)

    @app.post("/scout")
    async def run_scout(request: Request):
        """Run the job scout and add new relevant jobs to the tracker."""
        from datetime import datetime
        from ..discovery.scout import scout_jobs

        try:
            jobs = await asyncio.to_thread(scout_jobs)
        except Exception as e:
            # Re-render dashboard with error banner
            all_apps = load_applications()
            statuses = ["all", "discovered", "generated", "applied", "interviewing", "offered", "rejected", "not_relevant", "withdrawn"]
            stats = {s: sum(1 for a in all_apps if a.status == s) for s in statuses[1:]}
            stats["all"] = len(all_apps)
            companies = sorted({a.company for a in all_apps if a.company})
            sources = sorted({a.source for a in all_apps if a.source})
            return templates.TemplateResponse("dashboard.html", {
                "request": request,
                "applications": all_apps,
                "current_status": "all",
                "statuses": statuses,
                "stats": stats,
                "companies": companies,
                "sources": sources,
                "filters": {"q": "", "company": "", "source": "", "remote": ""},
                "scout_error": str(e),
            })

        added = 0
        for job in jobs:
            existing = get_application(job.id)
            if existing:
                continue
            app_record = Application(
                id=job.id,
                job_title=job.title,
                company=job.company,
                location=job.location,
                url=job.url,
                source=job.source,
                date_posted=job.date_posted,
                date_generated=datetime.now().isoformat(),
                status="discovered",
                description=job.description,
            )
            save_application(app_record)
            added += 1

        return RedirectResponse(f"/?status=discovered&scout_added={added}", status_code=302)

    # ── Profile routes ──────────────────────────────────────────────

    PROFILE_PATH = os.path.abspath(os.path.join(BASE_DIR, "config", "profile.yaml"))

    @app.get("/profile", response_class=HTMLResponse)
    async def profile_page(request: Request, saved: str = ""):
        try:
            profile = await asyncio.to_thread(load_profile, PROFILE_PATH)
        except Exception:
            profile = None  # No profile yet — show import section
        return templates.TemplateResponse("profile.html", {
            "request": request,
            "profile": profile,
            "save_ok": saved == "1",
            "save_error": None,
            "rebuild_error": None,
        })

    @app.post("/profile/save", response_class=HTMLResponse)
    async def profile_save(request: Request):
        import yaml

        form = await request.form()

        try:
            # ── Basic fields ──
            data: dict = {
                "name": form.get("name", "").strip(),
                "email": form.get("email", "").strip(),
                "phone": form.get("phone", "").strip(),
                "location": form.get("location", "").strip(),
                "citizenship": form.get("citizenship", "").strip(),
                "linkedin_url": form.get("linkedin_url", "").strip(),
                "summary": form.get("summary", "").strip(),
            }

            # ── Skills (dynamic categories) ──
            skills: dict[str, list[str]] = {}
            cat_idx = 0
            while True:
                cat_key = f"skill_category_{cat_idx}"
                if cat_key not in form:
                    cat_idx += 1
                    # check a few more in case indices are sparse (categories removed)
                    if cat_idx > 100:
                        break
                    continue
                cat_name = form[cat_key].strip()
                if not cat_name:
                    cat_idx += 1
                    continue
                skill_values = form.getlist(f"skill_{cat_idx}[]")
                skills[cat_name] = [s.strip() for s in skill_values if s.strip()]
                cat_idx += 1
            data["skills"] = skills

            # ── Experience ──
            titles = form.getlist("exp_title[]")
            companies = form.getlist("exp_company[]")
            types = form.getlist("exp_type[]")
            locations = form.getlist("exp_location[]")
            starts = form.getlist("exp_start[]")
            ends = form.getlist("exp_end[]")

            experiences = []
            for i in range(len(titles)):
                if not titles[i].strip():
                    continue
                bullets_raw = form.getlist(f"exp_bullet_{i}[]")
                bullets = [b.strip() for b in bullets_raw if b.strip()]
                exp = {
                    "title": titles[i].strip(),
                    "company": companies[i].strip() if i < len(companies) else "",
                    "type": types[i].strip() if i < len(types) else "",
                    "start": starts[i].strip() if i < len(starts) else "",
                    "end": ends[i].strip() if i < len(ends) else "",
                    "location": locations[i].strip() if i < len(locations) else "",
                    "bullets": bullets,
                }
                experiences.append(exp)
            data["experience"] = experiences

            # ── Education ──
            institutions = form.getlist("edu_institution[]")
            degrees = form.getlist("edu_degree[]")
            fields = form.getlist("edu_field[]")
            years = form.getlist("edu_years[]")

            educations = []
            for i in range(len(institutions)):
                if not institutions[i].strip():
                    continue
                edu = {
                    "institution": institutions[i].strip(),
                    "degree": degrees[i].strip() if i < len(degrees) else "",
                    "field": fields[i].strip() if i < len(fields) else "",
                    "years": years[i].strip() if i < len(years) else "",
                }
                educations.append(edu)
            data["education"] = educations

            # ── Simple lists ──
            data["certifications"] = [c.strip() for c in form.getlist("certification[]") if c.strip()]
            data["languages"] = [l.strip() for l in form.getlist("language[]") if l.strip()]
            data["project_highlights"] = [p.strip() for p in form.getlist("project_highlight[]") if p.strip()]

            # ── Write YAML ──
            def _write():
                with open(PROFILE_PATH, "w") as f:
                    yaml.dump(data, f, default_flow_style=False, allow_unicode=True, sort_keys=False)

            await asyncio.to_thread(_write)

            return RedirectResponse("/profile?saved=1", status_code=302)

        except Exception as exc:
            # On error, reload profile and show error
            profile = await asyncio.to_thread(load_profile, PROFILE_PATH)
            return templates.TemplateResponse("profile.html", {
                "request": request,
                "profile": profile,
                "save_ok": False,
                "save_error": str(exc),
            })

    @app.post("/profile/import", response_class=HTMLResponse)
    async def profile_import(request: Request):
        """Import profile from LinkedIn URL, resume PDF, or plain text."""
        from ..profile_builder import (
            build_profile_from_text,
            extract_text_from_docx,
            extract_text_from_pdf,
            fetch_linkedin_profile_text,
            save_profile,
        )

        form = await request.form()
        source_type = form.get("source_type", "")

        try:
            if source_type == "linkedin":
                url = form.get("linkedin_url", "").strip()
                if not url:
                    raise ValueError("Please provide a LinkedIn URL")
                text = await asyncio.to_thread(fetch_linkedin_profile_text, url)
                # Preserve the LinkedIn URL in the output
                profile_data = await asyncio.to_thread(build_profile_from_text, text)
                if not profile_data.get("linkedin_url"):
                    profile_data["linkedin_url"] = url

            elif source_type == "pdf":
                upload = form.get("resume_file")
                if not upload or not upload.filename:
                    raise ValueError("Please upload a file")
                file_bytes = await upload.read()
                filename = upload.filename.lower()
                if filename.endswith(".pdf"):
                    text = await asyncio.to_thread(extract_text_from_pdf, file_bytes)
                elif filename.endswith((".docx", ".doc")):
                    text = await asyncio.to_thread(extract_text_from_docx, file_bytes)
                elif filename.endswith(".txt"):
                    text = file_bytes.decode("utf-8", errors="replace")
                else:
                    raise ValueError(f"Unsupported file type: {upload.filename}")
                profile_data = await asyncio.to_thread(build_profile_from_text, text)

            elif source_type == "text":
                text = form.get("profile_text", "").strip()
                if not text or len(text) < 50:
                    raise ValueError("Please paste more text (at least 50 characters)")
                profile_data = await asyncio.to_thread(build_profile_from_text, text)

            else:
                raise ValueError(f"Unknown source type: {source_type}")

            await asyncio.to_thread(save_profile, profile_data)
            return RedirectResponse("/profile?saved=1", status_code=302)

        except Exception as exc:
            try:
                profile = await asyncio.to_thread(load_profile, PROFILE_PATH)
            except Exception:
                profile = None
            return templates.TemplateResponse("profile.html", {
                "request": request,
                "profile": profile,
                "save_ok": False,
                "save_error": None,
                "rebuild_error": str(exc),
            })

    # ── Settings routes ─────────────────────────────────────────────

    SETTINGS_PATH = os.path.abspath(os.path.join(BASE_DIR, "config", "settings.yaml"))

    @app.get("/settings", response_class=HTMLResponse)
    async def settings_page(request: Request, saved: str = ""):
        import yaml
        def _load():
            with open(SETTINGS_PATH, "r") as f:
                return yaml.safe_load(f)
        try:
            settings = await asyncio.to_thread(_load)
        except Exception:
            settings = {"scout": {}, "discovery": {}}
        # Ensure expected sections exist
        settings.setdefault("scout", {})
        settings.setdefault("discovery", {})
        return templates.TemplateResponse("settings.html", {
            "request": request,
            "settings": settings,
            "save_ok": saved == "1",
            "save_error": None,
        })

    @app.post("/settings/save", response_class=HTMLResponse)
    async def settings_save(request: Request):
        import yaml

        form = await request.form()

        try:
            # Load existing settings to preserve non-UI fields (llm, telegram, linkedin, output)
            def _load():
                with open(SETTINGS_PATH, "r") as f:
                    return yaml.safe_load(f) or {}
            settings = await asyncio.to_thread(_load)

            # ── Discovery ──
            settings.setdefault("discovery", {})
            settings["discovery"]["default_location"] = form.get("default_location", "").strip()
            settings["discovery"]["default_country"] = form.get("default_country", "").strip()

            # ── Scout ──
            settings.setdefault("scout", {})
            queries = [q.strip() for q in form.getlist("query[]") if q.strip()]
            settings["scout"]["queries"] = queries
            settings["scout"]["location"] = form.get("location", "").strip()
            settings["scout"]["country"] = form.get("country", "").strip()
            settings["scout"]["sources"] = form.getlist("source[]")
            settings["scout"]["max_per_query"] = int(form.get("max_per_query", 50))
            settings["scout"]["hours_old"] = int(form.get("hours_old", 336))
            settings["scout"]["interval_hours"] = int(form.get("interval_hours", 6))
            settings["scout"]["remote_only"] = form.get("remote_only") == "true"
            settings["scout"]["ai_filter"] = form.get("ai_filter") == "true"
            settings["scout"]["target_roles"] = form.get("target_roles", "").strip()

            # Title exclude keywords
            title_exclude = [kw.strip() for kw in form.getlist("title_exclude[]") if kw.strip()]
            settings["scout"]["title_exclude"] = title_exclude

            def _write():
                with open(SETTINGS_PATH, "w") as f:
                    yaml.dump(settings, f, default_flow_style=False, allow_unicode=True, sort_keys=False, width=120)

            await asyncio.to_thread(_write)
            return RedirectResponse("/settings?saved=1", status_code=302)

        except Exception as exc:
            import yaml
            try:
                def _reload():
                    with open(SETTINGS_PATH, "r") as f:
                        return yaml.safe_load(f) or {}
                settings = await asyncio.to_thread(_reload)
            except Exception:
                settings = {"scout": {}, "discovery": {}}
            settings.setdefault("scout", {})
            settings.setdefault("discovery", {})
            return templates.TemplateResponse("settings.html", {
                "request": request,
                "settings": settings,
                "save_ok": False,
                "save_error": str(exc),
            })

    @app.get("/view/{filename:path}")
    async def view_file(filename: str):
        """Extract and return formatted HTML content from a DOCX or text file."""
        if os.path.isabs(filename) and os.path.exists(filename):
            path = filename
        else:
            path = os.path.abspath(os.path.join(BASE_DIR, filename))

        if not os.path.exists(path):
            return JSONResponse({"html": "<p>File not found.</p>", "text": "File not found."}, status_code=404)

        if path.endswith(".docx"):
            html, text = _docx_to_html(path)
        elif path.endswith(".txt"):
            with open(path, "r") as f:
                text = f.read()
            import html as html_mod
            html = f"<pre style='white-space:pre-wrap;font-family:inherit;'>{html_mod.escape(text)}</pre>"
        elif path.endswith(".pdf"):
            html = "<p>(PDF preview not supported — use the download link.)</p>"
            text = "(PDF preview not supported — use the download link.)"
        else:
            html = "<p>(Unsupported file format.)</p>"
            text = "(Unsupported file format.)"

        return JSONResponse({"html": html, "text": text})

    def _docx_to_html(path: str) -> tuple[str, str]:
        """Convert a DOCX file to styled HTML and plain text."""
        import html as html_mod
        from docx import Document
        from docx.shared import Pt

        doc = Document(path)
        html_parts = []
        text_parts = []

        for para in doc.paragraphs:
            raw = para.text.strip()
            if not raw:
                continue

            text_parts.append(raw)

            # Detect formatting from runs
            is_all_bold = all(r.bold for r in para.runs if r.text.strip()) if para.runs else False
            font_size = None
            for r in para.runs:
                if r.font.size:
                    font_size = r.font.size
                    break

            escaped = html_mod.escape(raw)

            # Large bold = name header (22pt)
            if is_all_bold and font_size and font_size >= Pt(18):
                html_parts.append(f'<div style="font-size:1.6rem;font-weight:700;letter-spacing:0.5px;margin-bottom:2px;">{escaped}</div>')
            # Contact line (small, gray)
            elif font_size and font_size <= Pt(10) and "|" in raw:
                html_parts.append(f'<div style="font-size:0.85rem;color:#555;margin-bottom:12px;">{escaped}</div>')
            # Section headings (bold uppercase with border)
            elif is_all_bold and raw == raw.upper() and len(raw) > 3 and not raw.startswith("\u2022"):
                html_parts.append(
                    f'<div style="font-size:0.9rem;font-weight:700;text-transform:uppercase;'
                    f'border-bottom:1.5px solid #1a1a2e;padding-bottom:4px;margin:16px 0 8px;color:#1a1a2e;">{escaped}</div>'
                )
            # Job title lines (bold, with em-dash)
            elif is_all_bold and ("\u2014" in raw or " — " in raw):
                html_parts.append(f'<div style="font-weight:600;margin-top:10px;margin-bottom:2px;">{escaped}</div>')
            # Mixed bold/normal (e.g. "Category: skill, skill")
            elif any(r.bold for r in para.runs if r.text.strip()) and not all(r.bold for r in para.runs if r.text.strip()):
                parts = []
                for r in para.runs:
                    t = html_mod.escape(r.text)
                    if r.bold:
                        parts.append(f"<strong>{t}</strong>")
                    else:
                        parts.append(t)
                html_parts.append(f'<div style="margin-bottom:3px;">{"".join(parts)}</div>')
            # Bullet points
            elif raw.startswith("\u2022"):
                bullet_text = escaped.lstrip("\u2022").strip()
                html_parts.append(
                    f'<div style="padding-left:20px;margin-bottom:3px;position:relative;">'
                    f'<span style="position:absolute;left:4px;">&bull;</span>{bullet_text}</div>'
                )
            # Bold paragraph (e.g. sign-off name)
            elif is_all_bold:
                html_parts.append(f'<div style="font-weight:600;margin:4px 0;">{escaped}</div>')
            # Regular paragraph
            else:
                html_parts.append(f'<div style="margin-bottom:6px;">{escaped}</div>')

        return "\n".join(html_parts), "\n".join(text_parts)

    @app.get("/download/{filename:path}")
    async def download_file(filename: str):
        # Try absolute path first, then relative to project root
        if os.path.isabs(filename) and os.path.exists(filename):
            path = filename
        else:
            path = os.path.abspath(os.path.join(BASE_DIR, filename))

        if not os.path.exists(path):
            return HTMLResponse("File not found", status_code=404)
        return FileResponse(path, filename=os.path.basename(path))

    def _resolve_pdf_path(filename: str) -> str | None:
        """Resolve filename to an absolute PDF path, converting DOCX if needed."""
        if os.path.isabs(filename) and os.path.exists(filename):
            path = filename
        else:
            path = os.path.abspath(os.path.join(BASE_DIR, filename))

        if not os.path.exists(path):
            return None

        if path.endswith(".pdf"):
            return path

        if path.endswith(".docx"):
            pdf_path = path.replace(".docx", ".pdf")
            if not os.path.exists(pdf_path):
                pdf_path = _docx_to_pdf(path)
            return pdf_path

        return None

    @app.get("/pdf/{filename:path}")
    async def download_as_pdf(filename: str):
        """Serve file as PDF download — converts DOCX to PDF on the fly if needed."""
        path = _resolve_pdf_path(filename)
        if not path:
            return HTMLResponse("File not found", status_code=404)
        stem = os.path.splitext(os.path.basename(path))[0]
        return FileResponse(path, filename=f"{stem}.pdf", media_type="application/pdf")

    @app.get("/pdf-view/{filename:path}")
    async def view_pdf_inline(filename: str):
        """Serve PDF inline for iframe embedding (no Content-Disposition attachment)."""
        from starlette.responses import Response
        path = _resolve_pdf_path(filename)
        if not path:
            return HTMLResponse("File not found", status_code=404)
        with open(path, "rb") as f:
            content = f.read()
        return Response(content, media_type="application/pdf", headers={
            "Content-Disposition": "inline",
        })

    def _docx_to_pdf(docx_path: str) -> str:
        """Convert a DOCX file to PDF using WeasyPrint via an HTML intermediate."""
        import html as html_mod
        from docx import Document
        from weasyprint import HTML as WP_HTML

        doc = Document(docx_path)
        paragraphs_html = []

        for para in doc.paragraphs:
            raw = para.text
            if not raw.strip():
                paragraphs_html.append('<p style="margin:4px 0;">&nbsp;</p>')
                continue

            escaped = html_mod.escape(raw)
            is_all_bold = all(r.bold for r in para.runs if r.text.strip()) if para.runs else False
            font_size = None
            for r in para.runs:
                if r.font.size:
                    font_size = r.font.size
                    break

            from docx.shared import Pt

            if is_all_bold and font_size and font_size >= Pt(18):
                paragraphs_html.append(f'<p style="font-size:22pt;font-weight:700;margin:0 0 2px;">{escaped}</p>')
            elif font_size and font_size and font_size <= Pt(10) and "|" in raw:
                paragraphs_html.append(f'<p style="font-size:9pt;color:#444;margin:0 0 10px;">{escaped}</p>')
            elif is_all_bold and raw.strip() == raw.strip().upper() and len(raw.strip()) > 3:
                paragraphs_html.append(
                    f'<p style="font-size:10pt;font-weight:700;text-transform:uppercase;'
                    f'border-bottom:1px solid #000;padding-bottom:2px;margin:12px 0 5px;">{escaped}</p>'
                )
            elif is_all_bold and ("—" in raw or " — " in raw):
                paragraphs_html.append(f'<p style="font-size:10.5pt;font-weight:700;margin:8px 0 1px;">{escaped}</p>')
            elif raw.strip().startswith("•"):
                bullet = html_mod.escape(raw.strip().lstrip("•").strip())
                paragraphs_html.append(f'<p style="font-size:10pt;margin:1px 0 1px 18px;text-indent:-12px;">• {bullet}</p>')
            elif any(r.bold for r in para.runs if r.text.strip()):
                parts = []
                for r in para.runs:
                    t = html_mod.escape(r.text)
                    parts.append(f"<strong>{t}</strong>" if r.bold else t)
                paragraphs_html.append(f'<p style="font-size:10pt;margin:1px 0;">{"".join(parts)}</p>')
            else:
                paragraphs_html.append(f'<p style="font-size:10pt;margin:4px 0;">{escaped}</p>')

        html_content = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
  @page {{ margin: 0.6in 0.75in; }}
  body {{ font-family: Calibri, Arial, sans-serif; font-size: 10pt; color: #000; line-height: 1.3; }}
  p {{ margin: 0; }}
</style>
</head><body>{"".join(paragraphs_html)}</body></html>"""

        pdf_path = docx_path.replace(".docx", ".pdf")
        WP_HTML(string=html_content).write_pdf(pdf_path)
        return pdf_path

    return app
