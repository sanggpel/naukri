"""Render generated resumes and cover letters into PDF/DOCX files.

Matches the visual style of Sangeeta's existing resume:
- Clean, black text on white
- Large bold name, pipe-separated contact line underneath
- Bold uppercase section headings with a thin bottom border
- Bold job titles with em-dash: "Title — Company | Period | Location"
- Tight, compact spacing throughout
- Calibri/Arial font family
"""

import os
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from jinja2 import Environment, FileSystemLoader

from ..models import GeneratedResume


TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "templates")
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "data")

# Consistent font settings
FONT_NAME = "Calibri"
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)


def _get_jinja_env() -> Environment:
    return Environment(loader=FileSystemLoader(TEMPLATES_DIR))


def _has_weasyprint() -> bool:
    try:
        from weasyprint import HTML  # noqa: F401
        return True
    except ImportError:
        return False


def _set_font(run, size_pt, bold=False, color=BLACK, name=FONT_NAME):
    """Apply consistent font styling to a run."""
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = color
    # Also set East Asian font name to prevent fallback
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def _add_bottom_border(paragraph):
    """Add a thin black bottom border to a paragraph (like a horizontal rule)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")  # 0.5pt line
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_paragraph_spacing(p, before=0, after=0, line=None):
    """Set exact paragraph spacing in points."""
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if line is not None:
        fmt.line_spacing = Pt(line)


# ── Resume rendering ──────────────────────────────────────────────


def render_resume_pdf(
    resume: GeneratedResume,
    candidate_name: str = "Sangeeta Bahri",
    contact_info: dict = None,
) -> str:
    """Render resume. Uses WeasyPrint PDF if available, otherwise DOCX."""
    if _has_weasyprint():
        return _render_resume_weasyprint(resume, candidate_name, contact_info)
    else:
        return render_resume_docx(resume, candidate_name, contact_info)


def _render_resume_weasyprint(resume, candidate_name, contact_info=None):
    from weasyprint import HTML
    env = _get_jinja_env()
    template = env.get_template("resume_template.html")
    if contact_info is None:
        contact_info = _default_contact()
    html = template.render(
        name=candidate_name, contact=contact_info,
        resume=resume.sections, ats_score=resume.ats_score_estimate,
    )
    output_dir = os.path.join(OUTPUT_DIR, "resumes")
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    pdf_path = os.path.join(output_dir, f"resume_{timestamp}.pdf")
    HTML(string=html).write_pdf(pdf_path)
    return pdf_path


def render_resume_docx(
    resume: GeneratedResume,
    candidate_name: str = "Sangeeta Bahri",
    contact_info: dict = None,
) -> str:
    """Render resume to a clean, professional DOCX matching Sangeeta's style."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    if contact_info is None:
        contact_info = _default_contact()

    # ── NAME ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_spacing(p, before=0, after=0)
    run = p.add_run(candidate_name.upper())
    _set_font(run, size_pt=22, bold=True)

    # ── CONTACT LINE ──
    contact_parts = [
        contact_info["location"],
        contact_info["phone"],
        contact_info["email"],
        contact_info["linkedin"],
    ]
    if contact_info.get("citizenship"):
        contact_parts.append(contact_info["citizenship"])
    contact_str = " | ".join(contact_parts)

    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=0, after=6)
    run = p.add_run(contact_str)
    _set_font(run, size_pt=9.5, color=DARK_GRAY)

    # ── EXECUTIVE SUMMARY ──
    _add_resume_heading(doc, "EXECUTIVE SUMMARY")
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=0, after=4, line=13)
    run = p.add_run(resume.sections.executive_summary)
    _set_font(run, size_pt=10)

    # ── CORE COMPETENCIES ──
    _add_resume_heading(doc, "CORE COMPETENCIES")
    for category, skills in resume.sections.core_competencies.items():
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=0, after=1, line=13)
        run_cat = p.add_run(f"{category}: ")
        _set_font(run_cat, size_pt=10, bold=True)
        run_skills = p.add_run(", ".join(skills))
        _set_font(run_skills, size_pt=10)

    # ── PROFESSIONAL EXPERIENCE ──
    _add_resume_heading(doc, "PROFESSIONAL EXPERIENCE")
    for i, exp in enumerate(resume.sections.experience):
        # "Title — Company | Period | Location"
        p = doc.add_paragraph()
        top_space = 8 if i > 0 else 2
        _set_paragraph_spacing(p, before=top_space, after=1)

        title_str = f"{exp['title']} \u2014 {exp['company']}"
        meta_parts = []
        if exp.get("period"):
            meta_parts.append(exp["period"])
        if exp.get("location"):
            meta_parts.append(exp["location"])
        meta_str = " | ".join(meta_parts)

        run_title = p.add_run(title_str)
        _set_font(run_title, size_pt=10.5, bold=True)

        if meta_str:
            run_sep = p.add_run(f" | {meta_str}")
            _set_font(run_sep, size_pt=10)

        # Bullet points
        for bullet in exp.get("bullets", []):
            bp = doc.add_paragraph()
            _set_paragraph_spacing(bp, before=0, after=1, line=13)
            # Indent + bullet character
            bp.paragraph_format.left_indent = Inches(0.25)
            bp.paragraph_format.first_line_indent = Inches(-0.2)
            run = bp.add_run(f"\u2022  {bullet}")
            _set_font(run, size_pt=10)

    # ── EDUCATION & CERTIFICATIONS ──
    _add_resume_heading(doc, "EDUCATION & CERTIFICATIONS")
    for edu in resume.sections.education:
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=0, after=1, line=13)
        run = p.add_run(edu)
        _set_font(run, size_pt=10)

    if resume.sections.certifications:
        cert_line = " | ".join(resume.sections.certifications)
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=2, after=1, line=13)
        run = p.add_run(cert_line)
        _set_font(run, size_pt=10)

    # ── PROJECT HIGHLIGHTS ── (optional)
    if resume.sections.project_highlights:
        _add_resume_heading(doc, "PROJECT HIGHLIGHTS")
        for highlight in resume.sections.project_highlights:
            bp = doc.add_paragraph()
            _set_paragraph_spacing(bp, before=0, after=1, line=13)
            bp.paragraph_format.left_indent = Inches(0.25)
            bp.paragraph_format.first_line_indent = Inches(-0.2)
            run = bp.add_run(f"\u2022  {highlight}")
            _set_font(run, size_pt=10)

    # Save
    output_dir = os.path.join(OUTPUT_DIR, "resumes")
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_path = os.path.join(output_dir, f"resume_{timestamp}.docx")
    doc.save(docx_path)
    return docx_path


def _add_resume_heading(doc, text):
    """Section heading: bold, uppercase, with thin bottom border line."""
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=10, after=3)
    run = p.add_run(text)
    _set_font(run, size_pt=11, bold=True)
    _add_bottom_border(p)


# ── Cover Letter rendering ────────────────────────────────────────


def render_cover_letter_pdf(
    cover_letter_text: str,
    candidate_name: str = "Sangeeta Bahri",
) -> str:
    """Render cover letter. Uses WeasyPrint PDF if available, otherwise DOCX."""
    if _has_weasyprint():
        return _render_cl_weasyprint(cover_letter_text, candidate_name)
    else:
        return render_cover_letter_docx(cover_letter_text, candidate_name)


def _render_cl_weasyprint(cover_letter_text, candidate_name):
    from weasyprint import HTML
    env = _get_jinja_env()
    template = env.get_template("cover_letter_template.html")
    html = template.render(name=candidate_name, content=cover_letter_text)
    output_dir = os.path.join(OUTPUT_DIR, "cover_letters")
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    pdf_path = os.path.join(output_dir, f"cover_letter_{timestamp}.pdf")
    HTML(string=html).write_pdf(pdf_path)
    return pdf_path


def render_cover_letter_docx(
    cover_letter_text: str,
    candidate_name: str = "Sangeeta Bahri",
) -> str:
    """Render cover letter to a clean, professional DOCX."""
    doc = Document()

    # Margins (standard letter margins)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Parse the cover letter text and render each paragraph
    lines = cover_letter_text.strip().split("\n")
    current_para = []

    for line in lines:
        stripped = line.strip()
        if stripped == "":
            # Flush current paragraph
            if current_para:
                _add_cl_paragraph(doc, " ".join(current_para))
                current_para = []
        else:
            # Check if this is a short standalone line (greeting, sign-off, contact)
            is_standalone = (
                stripped.startswith("Dear ")
                or stripped.startswith("Sincerely")
                or stripped.startswith("Best ")
                or stripped.startswith("Regards")
                or stripped == candidate_name
                or "@" in stripped  # contact line with email
                or (len(stripped) < 60 and any(c.isdigit() for c in stripped) and "(" in stripped)  # phone number line
            )

            if is_standalone:
                # Flush any pending paragraph first
                if current_para:
                    _add_cl_paragraph(doc, " ".join(current_para))
                    current_para = []
                _add_cl_paragraph(doc, stripped, space_after=2)
            else:
                current_para.append(stripped)

    # Flush remaining
    if current_para:
        _add_cl_paragraph(doc, " ".join(current_para))

    # Save
    output_dir = os.path.join(OUTPUT_DIR, "cover_letters")
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_path = os.path.join(output_dir, f"cover_letter_{timestamp}.docx")
    doc.save(docx_path)
    return docx_path


def _add_cl_paragraph(doc, text, space_after=8):
    """Add a cover letter paragraph with consistent styling."""
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=0, after=space_after, line=15)
    run = p.add_run(text)
    _set_font(run, size_pt=11)


# ── Helpers ───────────────────────────────────────────────────────


def _default_contact() -> dict:
    """Load contact info from profile.yaml. Falls back to empty strings."""
    try:
        from ..profile_loader import load_profile
        profile = load_profile()
        linkedin_short = profile.linkedin_url.replace("https://", "").replace("http://", "").rstrip("/")
        return {
            "phone": profile.phone,
            "email": profile.email,
            "linkedin": linkedin_short,
            "location": profile.location,
            "citizenship": getattr(profile, "citizenship", ""),
        }
    except Exception:
        return {"phone": "", "email": "", "linkedin": "", "location": "", "citizenship": ""}
